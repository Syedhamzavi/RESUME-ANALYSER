from typing import Dict, List
import re
from statistics import mean
from collections import defaultdict
from docx import Document
from PyPDF2 import PdfReader
import pdfplumber
import io

FONT_SIZE_SMALL_PENALTY = 12
FONT_SIZE_LARGE_PENALTY = 6
NO_INDENTATION_PENALTY = 6
MISSING_EMAIL_PENALTY = 10
MISSING_PHONE_PENALTY = 8
MISSING_LINKEDIN_PENALTY = 5
MISSING_GITHUB_PENALTY = 4
MISSING_SECTION_PENALTY = 15

SKILLS_DATABASE = {
    "programming": ["python", "java", "javascript", "c++", "c#", "ruby", "php", "swift", "kotlin", "go", "rust", "typescript", "html", "css", "sql", "r", "matlab", "scala", "perl", "bash", "shell", "powershell", "dart", "objective-c", "assembly", "fortran", "cobol", "lua"],
    "frameworks": ["django", "flask", "spring", "angular", "vue", "express", "laravel", "rails", "asp.net", "tensorflow", "pytorch", "keras", "node.js", "react", "react native", "flutter", "jquery", "bootstrap", "ember", "backbone", "meteor", "svelte", "next.js", "nuxt.js", "nestjs", "fastapi", "graphql", "hibernate", "mybatis", "jpa"],
    "databases": ["mysql", "postgresql", "mongodb", "redis", "oracle", "sql server", "cassandra", "elasticsearch", "dynamodb", "firebase", "cosmosdb", "sqlite", "mariadb", "couchdb", "neo4j", "arangodb", "rethinkdb", "couchbase", "memcached", "hbase", "bigtable"],
    "cloud": ["aws", "azure", "google cloud", "docker", "kubernetes", "terraform", "ansible", "jenkins", "ci/cd", "serverless", "lambda", "ec2", "s3", "azure functions", "google functions", "cloud formation", "cloudwatch", "azure devops", "github actions", "circleci", "gitlab ci", "travis ci", "heroku", "digital ocean", "linode", "vultr"],
    "tools": ["git", "github", "gitlab", "jira", "confluence", "slack", "trello", "jenkins", "circleci", "github actions", "docker", "kubernetes", "postman", "swagger", "visual studio", "intellij", "eclipse", "vs code", "android studio", "xcode", "webstorm", "pycharm", "phpstorm", "rubymine", "sublime", "atom", "notepad++"],
    "methodologies": ["agile", "scrum", "kanban", "waterfall", "devops", "ci/cd", "tdd", "bdd", "pair programming", "code review", "version control", "microservices", "monolith", "rest", "soap", "graphql", "grpc", "oauth", "jwt", "openid", "saml"],
    "soft_skills": ["leadership", "communication", "teamwork", "problem solving", "critical thinking", "adaptability", "time management", "creativity", "collaboration", "presentation", "mentoring", "coaching", "negotiation", "conflict resolution", "decision making", "strategic thinking", "analytical skills", "attention to detail", "multitasking"]
}

all_skills = [skill for category_skills in SKILLS_DATABASE.values() for skill in category_skills]
skill_to_category_map = {skill: cat for cat, skills in SKILLS_DATABASE.items() for skill in skills}
sorted_skills = sorted(all_skills, key=len, reverse=True)
skills_pattern = r'\b(' + '|'.join(re.escape(skill) for skill in sorted_skills) + r')\b'
COMPILED_SKILLS_REGEX = re.compile(skills_pattern, re.IGNORECASE)

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+91[-\s]*)?(?:\d{3}[-\s]*\d{3}[-\s]*\d{4}|\d{10})")
LINKEDIN_RE = re.compile(r"linkedin\.com/(in|pub)/[a-zA-Z0-9_-]+", re.I)
GITHUB_RE = re.compile(r"github\.com/[a-zA-Z0-9_.-]{1,39}/?", re.I)

SECTION_KEYWORDS = {
    "summary": ["summary", "objective", "profile"],
    "education": ["education", "academic", "qualification", "coursework"],
    "experience": ["experience", "employment", "work history", "professional"],
    "skills": ["skills", "technical skills", "technologies", "tools", "tech stack"],
    "certificates": ["certificates", "certifications", "certification", "licenses", "courses"],
    "projects": ["projects", "academic projects", "personal projects"],
}

def extract_skills(text: str) -> Dict[str, List[str]]:
    if not text or not isinstance(text, str): return {}
    found_skills = defaultdict(set)
    try:
        matches = COMPILED_SKILLS_REGEX.finditer(text)
        for match in matches:
            skill = match.group(0).lower()
            category = skill_to_category_map.get(skill)
            if category:
                found_skills[category].add(skill)
        return {cat: sorted(list(skills)) for cat, skills in found_skills.items()}
    except Exception:
        return {}

def extract_experience(text: str) -> int:
    if not text or not isinstance(text, str): return 0
    max_experience = 0
    try:
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of)?\s*experience', r'experience\s*:\s*(\d+)\+?\s*years?',
            r'(\d+)\s*years?\s*(?:in|of)', r'(\d+)\s*-\s*(\d+)\s*years?\s*experience',
            r'minimum\s*of\s*(\d+)\s*years', r'at least\s*(\d+)\s*years'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    exp = max(int(m) for m in match if m.isdigit()) if isinstance(match, tuple) else int(match)
                    max_experience = max(max_experience, exp)
                except ValueError:
                    continue
    except Exception: return 0
    return max_experience

def extract_education_from_text(text: str) -> List[Dict]:
    education_patterns = [
        (r'\b(bachelor\s+of\s+engineering|b\.?e\.?)\b', 'bachelor'),(r'\b(bachelor\s+of\s+technology|b\.?tech)\b', 'bachelor'),
        (r'\b(bachelor\s+of\s+science|b\.?sc\.?)\b', 'bachelor'), (r'\b(bachelor[\'s]?)\b', 'bachelor'),
        (r'\b(master\s+of\s+engineering|m\.?e\.?)\b', 'master'), (r'\b(master\s+of\s+technology|m\.?tech)\b', 'master'),
        (r'\b(master\s+of\s+science|m\.?sc\.?)\b', 'master'), (r'\b(master\s+of\s+business\s+administration|mba)\b', 'master'),
        (r'\b(master[\'s]?)\b', 'master'), (r'\b(doctorate|phd|ph\.d)\b', 'phd'),
    ]
    found_education, seen_education = [], set()
    sentences = re.split(r'[.\n]', text)
    for sentence in sentences:
        if not sentence.strip(): continue
        for pattern, degree_type in education_patterns:
            match = re.search(pattern, sentence, re.IGNORECASE)
            if match:
                remaining = sentence[match.end():]
                field_match = re.search(r'^\s*(?:in|of)?\s+([\w\s]+(?:(?:and|&)\s*[\w\s]+)*)', remaining, re.IGNORECASE)
                field = field_match.group(1).strip() if field_match else ""
                edu_key = (degree_type, field.lower())
                if edu_key not in seen_education:
                    found_education.append({"degree": degree_type, "field": field})
                    seen_education.add(edu_key)
                break
    return found_education

def _extract_docx(bytes_io):
    try:
        doc = Document(bytes_io)
        text_lines, font_sizes, line_spacings, indent_count = [], [], [], 0
        for para in doc.paragraphs:
            txt = (para.text or "").strip()
            if txt: text_lines.append(txt)
            for run in para.runs:
                if run.font.size: font_sizes.append(run.font.size.pt)
            if para.paragraph_format.line_spacing:
                try: line_spacings.append(float(para.paragraph_format.line_spacing))
                except (ValueError, TypeError): pass
            if (para.paragraph_format.left_indent and getattr(para.paragraph_format.left_indent, "pt", 0) > 2) or txt.startswith((" ", "\t", "•", "-", "*")):
                indent_count += 1
        word_count = len(" ".join(text_lines).split())
        page_count = len(doc.paragraphs) // 30 or 1
        return "\n".join(text_lines), font_sizes, line_spacings, indent_count, word_count, page_count
    except Exception:
        return "", [], [], 0, 0, 0

def _extract_pdf(bytes_io):
    text_lines, font_sizes, line_spacings, indent_count, page_count = [], [], [], 0, 0
    full_text = ""
    try:
        bytes_io.seek(0)
        with pdfplumber.open(bytes_io) as pdf:
            page_count = len(pdf.pages)
            full_text = "".join([page.extract_text(x_tolerance=2) or "" for page in pdf.pages])
            for page in pdf.pages:
                for ch in page.chars:
                    if isinstance(ch.get("size"), (int, float)):
                        font_sizes.append(float(ch.get("size")))
    except Exception:
        try:
            bytes_io.seek(0)
            reader = PdfReader(bytes_io)
            page_count = len(reader.pages)
            full_text = "".join([page.extract_text() or "" for page in reader.pages])
        except Exception:
            pass
    if full_text:
        text_lines = [ln.strip("\r") for ln in full_text.split("\n") if ln.strip()]
    word_count = len(" ".join(text_lines).split())
    indent_count = sum(1 for ln in text_lines if ln.strip().startswith((" ", "\t", "•", "-", "*")))
    return "\n".join(text_lines), font_sizes, line_spacings, indent_count, word_count, max(1, page_count)

def _extract_contacts(text: str):
    def first(rx):
        m = rx.search(text)
        return m.group(0) if m else None
    
    raw_phone = first(PHONE_RE)
    if raw_phone and len(re.sub(r"\D", "", raw_phone)) < 10:
        raw_phone = None
    
    return {
        "email": first(EMAIL_RE), "phone": raw_phone,
        "linkedin": first(LINKEDIN_RE), "github": first(GITHUB_RE)
    }

def _detect_sections_presence(text: str):
    return {sec: any(k in text.lower() for k in keys) for sec, keys in SECTION_KEYWORDS.items()}

def _line_spacing_ratio(text: str):
    lines = text.splitlines()
    return round(sum(1 for l in lines if not l.strip()) / max(1, len(lines)), 2)

def analyze_resume_bytes(bytes_io, filename: str) -> Dict:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        text, font_sizes, line_spacings, indents, word_count, page_count = _extract_docx(bytes_io)
    elif name.endswith(".pdf"):
        text, font_sizes, line_spacings, indents, word_count, page_count = _extract_pdf(bytes_io)
    else:
        return {"error": "Unsupported file format. Upload PDF or DOCX."}

    try:
        text = re.sub(r'java\s+script', 'javascript', text, flags=re.IGNORECASE)
        text = re.sub(r'tensor\s*f\s*low', 'tensorflow', text, flags=re.IGNORECASE)
        text = re.sub(r'matplo\s*tlib', 'matplotlib', text, flags=re.IGNORECASE)
        text = re.sub(r'scikit\s*-\s*learn', 'scikit-learn', text, flags=re.IGNORECASE)
    except Exception:
        pass

    if not text or not text.strip():
        return {"error": "Could not extract text from the file."}

    contacts = _extract_contacts(text)
    sections = _detect_sections_presence(text)
    
    score, suggestions = 100, []
    avg_font_size = round(mean(font_sizes), 1) if font_sizes else 11.0
    if avg_font_size < 10:
        score -= FONT_SIZE_SMALL_PENALTY; suggestions.append("Font size is small. Use 11–12pt for readability.")
    elif avg_font_size > 13.5:
        score -= FONT_SIZE_LARGE_PENALTY; suggestions.append("Font size is large. Prefer an 11–12pt range.")
    
    if indents == 0:
        score -= NO_INDENTATION_PENALTY; suggestions.append("Use bullets/indentation to improve scannability.")

    if not contacts.get("email"):
        score -= MISSING_EMAIL_PENALTY; suggestions.append("Add a valid email address.")
    if not contacts.get("phone"):
        score -= MISSING_PHONE_PENALTY; suggestions.append("Add a reachable phone number.")
    if not contacts.get("linkedin"):
        score -= MISSING_LINKEDIN_PENALTY; suggestions.append("Consider adding a LinkedIn profile link.")

    is_tech = any(k in text.lower() for k in ['developer', 'engineer', 'scientist', 'data', 'software'])
    if not contacts.get("github") and is_tech:
        score -= MISSING_GITHUB_PENALTY
        suggestions.append("This appears to be a technical role; a GitHub profile is highly recommended.")

    missing_sections = [sec for sec in ["experience", "skills", "education"] if not sections.get(sec)]
    if missing_sections:
        score -= MISSING_SECTION_PENALTY
        suggestions.append(f"Add missing critical sections: {', '.join(missing_sections)}.")

    return {
        "score": max(0, score),
        "metrics": {
            "avg_font_size": avg_font_size, "avg_line_spacing": round(mean(line_spacings), 2) if line_spacings else _line_spacing_ratio(text),
            "indentations": indents, "blank_lines": sum(1 for l in text.splitlines() if not l.strip()),
            "sections_found": sections, "words": word_count, "pages": page_count,
        },
        "contacts": contacts,
        "suggestions": suggestions,
        "full_text": text,
        "education": extract_education_from_text(text),
        "experience": extract_experience(text),
        "skills": extract_skills(text)
    }